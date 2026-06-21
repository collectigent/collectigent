"""文档加载器 - 支持多种文件格式"""

from __future__ import annotations

from abc import ABC, abstractmethod
from enum import Enum
from typing import List, Optional, Any
import os


class FileType(Enum):
    """文件类型枚举"""
    TXT = "txt"
    MARKDOWN = "md"
    PDF = "pdf"
    DOCX = "docx"
    JSON = "json"
    HTML = "html"


class Document:
    """文档对象"""
    
    def __init__(
        self,
        content: str,
        source: str,
        file_type: FileType,
        metadata: Optional[dict] = None,
        id: Optional[str] = None,
    ):
        self.content = content
        self.source = source
        self.file_type = file_type
        self.metadata = metadata or {}
        self.id = id or self._generate_id()
    
    def _generate_id(self) -> str:
        """生成唯一ID"""
        import hashlib
        return hashlib.md5(f"{self.source}{self.content[:100]}".encode()).hexdigest()
    
    def __len__(self) -> int:
        return len(self.content)
    
    def __repr__(self) -> str:
        return f"Document(source={self.source}, length={len(self)})"


class DocumentLoader(ABC):
    """文档加载器基类"""
    
    @abstractmethod
    def load(self, path: str) -> Document:
        """加载单个文档"""
        pass
    
    @abstractmethod
    def load_dir(self, dir_path: str, patterns: Optional[List[str]] = None) -> List[Document]:
        """加载目录下的所有文档"""
        pass
    
    @staticmethod
    def detect_file_type(file_path: str) -> FileType:
        """检测文件类型"""
        ext = file_path.lower().split(".")[-1]
        mapping = {
            "txt": FileType.TXT,
            "md": FileType.MARKDOWN,
            "markdown": FileType.MARKDOWN,
            "pdf": FileType.PDF,
            "docx": FileType.DOCX,
            "json": FileType.JSON,
            "html": FileType.HTML,
        }
        return mapping.get(ext, FileType.TXT)
    
    @staticmethod
    def create(file_type: Optional[FileType] = None) -> "DocumentLoader":
        """创建对应的加载器"""
        if file_type in (FileType.PDF, None):
            return PDFLoader()
        elif file_type == FileType.DOCX:
            return DocxLoader()
        else:
            return TextLoader()


class TextLoader(DocumentLoader):
    """文本文件加载器（支持txt/md/json/html）"""
    
    def load(self, path: str) -> Document:
        """加载文本文件"""
        file_type = self.detect_file_type(path)
        
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        
        return Document(
            content=content,
            source=path,
            file_type=file_type,
            metadata={"file_size": len(content)},
        )
    
    def load_dir(self, dir_path: str, patterns: Optional[List[str]] = None) -> List[Document]:
        """加载目录下的文本文件"""
        documents = []
        patterns = patterns or ["*.txt", "*.md", "*.json", "*.html"]
        
        import glob
        
        for pattern in patterns:
            for file_path in glob.glob(os.path.join(dir_path, pattern)):
                if os.path.isfile(file_path):
                    try:
                        doc = self.load(file_path)
                        documents.append(doc)
                    except Exception as e:
                        print(f"Failed to load {file_path}: {e}")
        
        return documents


class PDFLoader(DocumentLoader):
    """PDF文件加载器"""
    
    def __init__(self):
        self._import_dependencies()
    
    def _import_dependencies(self):
        """导入PDF依赖"""
        try:
            global PyPDF2
            import PyPDF2
        except ImportError:
            raise ImportError("请安装PyPDF2: pip install PyPDF2")
    
    def load(self, path: str) -> Document:
        """加载PDF文件"""
        content = ""
        
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                content += page.extract_text() + "\n"
        
        return Document(
            content=content.strip(),
            source=path,
            file_type=FileType.PDF,
            metadata={"pages": len(reader.pages)},
        )
    
    def load_dir(self, dir_path: str, patterns: Optional[List[str]] = None) -> List[Document]:
        """加载目录下的PDF文件"""
        documents = []
        patterns = patterns or ["*.pdf"]
        
        import glob
        
        for pattern in patterns:
            for file_path in glob.glob(os.path.join(dir_path, pattern)):
                if os.path.isfile(file_path):
                    try:
                        doc = self.load(file_path)
                        documents.append(doc)
                    except Exception as e:
                        print(f"Failed to load {file_path}: {e}")
        
        return documents


class DocxLoader(DocumentLoader):
    """DOCX文件加载器"""
    
    def __init__(self):
        self._import_dependencies()
    
    def _import_dependencies(self):
        """导入DOCX依赖"""
        try:
            global DocxDocument
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装python-docx: pip install python-docx")
    
    def load(self, path: str) -> Document:
        """加载DOCX文件"""
        doc = DocxDocument(path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return Document(
            content=content.strip(),
            source=path,
            file_type=FileType.DOCX,
            metadata={"paragraphs": len(doc.paragraphs)},
        )
    
    def load_dir(self, dir_path: str, patterns: Optional[List[str]] = None) -> List[Document]:
        """加载目录下的DOCX文件"""
        documents = []
        patterns = patterns or ["*.docx"]
        
        import glob
        
        for pattern in patterns:
            for file_path in glob.glob(os.path.join(dir_path, pattern)):
                if os.path.isfile(file_path):
                    try:
                        doc = self.load(file_path)
                        documents.append(doc)
                    except Exception as e:
                        print(f"Failed to load {file_path}: {e}")
        
        return documents